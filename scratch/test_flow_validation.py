import sys
from pathlib import Path

# Configure stdout for utf-8 on Windows consoles
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

ROOT_DIR = Path(__file__).resolve().parent.parent
SRC_DIR = ROOT_DIR / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

import docx
from ingestion.document_processor import DocumentProcessor
from core.embedding_manager import EmbeddingManager
from retrieval.keyword_search import KeywordSearch
from retrieval.retriever import HybridRetriever
from llm.generator import LLMGenerator
from core.pipeline import RAGPipeline

def run_e2e_test():
    print("=" * 80)
    print("END-TO-END RAG FLOW VALIDATION: DOCX -> 14 CHUNKS -> RETRIEVAL -> GROQ")
    print("=" * 80)

    # STEP 1: CREATE DOCX FILE WITH 14 DISTINCT SECTIONS
    print("\n[STEP 1] Generating DOCX with 14 substantive sections...")
    docx_path = ROOT_DIR / "data" / "SDACOURSEMATERIAL.docx"
    doc = docx.Document()
    doc.add_heading("Streaming Data Analytics Course Material", 0)

    # 14 distinct topics with full paragraphs
    topics = [
        ("Section 1: Online Regression Architecture",
         "Online regression is a supervised learning methodology where linear models update their internal weight coefficients dynamically as each data point arrives sequentially over an infinite stream. Unlike traditional batch offline regression that minimizes total empirical loss across the entire dataset stored in persistent storage, online regression minimizes instantaneous loss per step using online gradient updates. This makes it exceptionally well-suited for high-velocity streaming architectures, low-latency financial feeds, and IoT edge analytics where memory constraints prevent holding historical batches."),
        
        ("Section 2: Online Gradient Descent Optimization",
         "Online gradient descent processes streaming records sequentially. For each arriving input vector x_t and ground truth label y_t, the loss gradient is computed and weights w_t are adjusted immediately using step size eta_t: w_{t+1} = w_t - eta_t * grad(loss). By applying dynamic step-size decay schedules such as 1/sqrt(t) or adaptive scaling, online gradient descent guarantees asymptotic convergence to the optimal parameter manifold while maintaining constant O(d) space complexity per iteration."),
        
        ("Section 3: Batch Learning vs Streaming Online Learning",
         "Batch learning requires all historical observations to reside simultaneously in memory to compute closed-form matrix inversions or global epoch iterations. In contrast, online streaming algorithms process each record in a single pass with strict bounded memory and update times. When streaming workloads encounter memory exhaustion or infinite data cardinality, online algorithms provide continuous availability and instantaneous parameter adaptation without expensive retraining."),
        
        ("Section 4: Concept Drift Detection Mechanisms",
         "Concept drift represents changes in the underlying statistical distribution P(X, Y) over time in non-stationary data streams. Drift detection algorithms such as DDM (Drift Detection Method), EDDM, and ADWIN (Adaptive Windowing) monitor tracking error rates and confidence intervals. When a statistically significant distribution shift is detected, older observations are pruned, model learning rates are dynamically reset, and ensembles adapt immediately to the newly emerging data pattern."),
        
        ("Section 5: Hoeffding Trees and Streaming Classifiers",
         "Hoeffding Trees, also known as Very Fast Decision Trees (VFDT), are incremental decision tree induction algorithms designed specifically for massive data streams. They use the statistical Hoeffding bound to determine the minimum number of streaming instances needed to select an optimal splitting attribute with high theoretical confidence, ensuring the streaming decision tree is asymptotically identical to a batch tree without holding all data in RAM."),
        
        ("Section 6: Stochastic Gradient Descent Variations",
         "Stochastic Gradient Descent (SGD) minimizes the expected risk function along noisy unbiased gradient estimates calculated from random single observations. In streaming environments, SGD eliminates offline shuffling requirements and provides optimal linear-time throughput. Extensions like Mini-batch SGD, momentum accelerators, and Nesterov accelerated gradients stabilize stochastic trajectory oscillations across steep loss ravines."),
        
        ("Section 7: Adaptive Learning Rate Algorithms",
         "Adaptive learning rate optimizers dynamically scale update step sizes on a per-parameter basis. Algorithms like AdaGrad accumulate past squared gradients to dampen frequent feature updates, while RMSprop and Adam introduce exponentially decaying running averages of past gradients and squared gradients. These mechanisms enable online regression models to self-tune learning rates across sparse and noisy streaming signals without manual hyperparameter interventions."),
        
        ("Section 8: Sliding Window Processing Models",
         "Sliding window processing models maintain a FIFO buffer of the most recent N streaming events, automatically discarding stale observations outside the active horizon. By bounding memory consumption and computational complexity to window size N, sliding windows naturally track evolving non-stationary data distributions and prevent historical data bias from dominating current model predictions."),
        
        ("Section 9: Reservoir Sampling on Unbounded Streams",
         "Reservoir sampling (Algorithm R) maintains an unbiased uniform random sample of size k from an unbounded continuous data stream of unknown total length N in a single linear pass. For the t-th incoming item, it is inserted into the reservoir with probability k/t, replacing a randomly chosen existing element. This guarantees equal sampling probability for all stream elements using strictly O(k) memory."),
        
        ("Section 10: Count-Min Sketch Frequency Estimation",
         "The Count-Min Sketch is a probabilistic sublinear space data structure that estimates frequency moments and item frequencies in massive streaming networks. Using an array of w counters across d independent pairwise hash functions, it provides point queries with provable error bounds (epsilon) with probability (1 - delta), requiring exponentially less memory than full hash maps."),
        
        ("Section 11: HyperLogLog Cardinality Estimation",
         "HyperLogLog (HLL) is an advanced streaming sketch algorithm that estimates the number of distinct elements (cardinality) in massive datasets with standard error approximately 1.04 / sqrt(m) using m memory registers. By observing the distribution of leading zeros in hashed stream identifiers, HLL accurately counts billions of unique stream events using less than 1.5 kilobytes of memory."),
        
        ("Section 12: Exponential Smoothing and Time Series Forecasting",
         "Exponential smoothing methods apply exponentially decreasing weights to historical stream observations, placing highest importance on the most recent data points while retaining an exponentially fading memory of past patterns. Holt-Winters extensions incorporate dynamic level, trend, and seasonal tracking components, delivering real-time anomaly detection and predictive stream forecasting."),
        
        ("Section 13: Incremental Matrix Factorization for Recommendations",
         "Incremental matrix factorization algorithms continuously update low-rank latent factor representations for users and items as streaming interaction events (clicks, purchases, ratings) occur. By updating latent vectors using stochastic gradient steps on incoming events, recommendation models reflect instant user preference changes without performing expensive offline matrix decompositions."),
        
        ("Section 14: Prequential Evaluation and Real-Time Testing",
         "Prequential evaluation (test-then-train) is the gold standard benchmark protocol for streaming machine learning models. Each streaming instance is first used to test and evaluate the current model state (generating an unbiased prediction and measuring loss) before being used to update the model parameters. This ensures continuous, real-time performance tracking without separate test splits.")
    ]

    for title, content in topics:
        doc.add_paragraph(f"{title}\n{content}")

    doc.save(str(docx_path))
    print(f"DOCX successfully created at: {docx_path}")

    # STEP 2: INGESTION & CHUNKING (VERIFY 14 CHUNKS)
    print("\n[STEP 2] Ingesting and chunking DOCX...")
    processor = DocumentProcessor(min_chunk_size=10, max_chunk_size=100, chunk_overlap=0)
    chunks = processor.process_document(str(docx_path), user_id="test_user_jwt_123")
    print(f"Total Chunks Created: {len(chunks)}")
    assert len(chunks) == 14, f"Expected 14 chunks, got {len(chunks)}"
    for idx, c in enumerate(chunks):
        print(f"  Chunk [{idx + 1:02d}] Source: {c['metadata']['source']} | User: {c['metadata'].get('user_id')} | Words: {len(c['text'].split()):02d} | Text: {c['text'][:65]}...")

    # STEP 3: EMBEDDINGS
    print("\n[STEP 3] Generating Embeddings via all-MiniLM-L6-v2...")
    emb_mgr = EmbeddingManager(model_name="all-MiniLM-L6-v2")
    emb_mgr.add_chunks(chunks, save=False)
    print(f"FAISS Index Count: {emb_mgr.index.ntotal} vectors | Dimension: {emb_mgr.dimension}")

    # STEP 4: VECTOR SEARCH (FAISS) & BM25 SEARCH
    print("\n[STEP 4] Vector Search (FAISS) with active session user_id...")
    kw_search = KeywordSearch(Path("embeddings/bm25_flow_test.pkl"))
    kw_search.add_chunks(chunks)

    test_query = "What is online regression and how does it update parameters dynamically?"
    print(f"Query: '{test_query}'")

    faiss_candidates = emb_mgr.search(test_query, top_k=5, user_id="test_user_jwt_123")
    print(f"FAISS Matches Found: {len(faiss_candidates)}")
    for i, c in enumerate(faiss_candidates[:2]):
        print(f"  FAISS Top [{i+1}] L2 Distance Score: {c['score']:.4f} | Snippet: {c['text'][:80]}...")

    # STEP 5: HYBRID RETRIEVAL & RELEVANT CHUNKS
    print("\n[STEP 5] Hybrid Retrieval (FAISS + BM25 Fusion)...")
    retriever = HybridRetriever(emb_mgr, kw_search)
    retrieved = retriever.retrieve(test_query, top_k=5, user_id="test_user_jwt_123")
    print(f"Retrieved Relevant Chunks Count: {len(retrieved)}")
    assert len(retrieved) > 0, "No chunks retrieved!"

    # STEP 6: RERANKING
    print("\n[STEP 6] Cross-Encoder Reranking...")
    for i, c in enumerate(retrieved):
        print(f"  Chunk [{i+1}] Raw Logit: {c.get('raw_rerank_score'):.4f} | Rerank Sigmoid: {c.get('rerank_score'):.6f} | Text: {c['text'][:80]}...")

    # STEP 7: CONFIDENCE GATE
    print("\n[STEP 7] Confidence Gate Verification...")
    llm = LLMGenerator()
    pipeline = RAGPipeline(retriever, llm, confidence_threshold=0.15)
    best_score = retrieved[0].get('rerank_score', 0.0)
    print(f"Best Score: {best_score:.6f} | Threshold: {pipeline.confidence_threshold}")
    print(f"Passes Confidence Threshold? {best_score >= pipeline.confidence_threshold}")
    assert best_score >= pipeline.confidence_threshold, f"Score {best_score} failed threshold {pipeline.confidence_threshold}"

    # STEP 8: CONTEXT PREPARATION
    print("\n[STEP 8] Context Preparation for LLM Generation...")
    print("Context blocks generated with citation headers:")
    for i, chunk in enumerate(retrieved[:2]):
        print(f"  --- Document [{i+1}] Source: {chunk['metadata']['source']} ---\n  {chunk['text'][:120]}...")

    # STEP 9 & 10: FULL PIPELINE EXECUTION (Groq + Actual Answer)
    print("\n[STEP 9 & 10] Executing Full Pipeline (Calling Groq LLM)...")
    answer, meta = pipeline.run(test_query, user_id="test_user_jwt_123", session_id="test_session")
    print("\n" + "=" * 80)
    print("FINAL ACTUAL ANSWER:")
    print("=" * 80)
    print(answer)
    print("=" * 80)
    print("EXECUTION METRICS:")
    print(f"Confidence Score: {meta.get('confidence'):.6f} ({meta.get('confidence') * 100:.2f}%)")
    print(f"Total Latency: {meta.get('latency')}s (Retrieval: {meta.get('retrieval_time')}s, LLM: {meta.get('llm_time')}s)")
    print(f"Sources Used: {len(meta.get('sources', []))}")
    print("=" * 80)

if __name__ == "__main__":
    run_e2e_test()
