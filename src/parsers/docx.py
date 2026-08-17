import io
from typing import List, Dict, Any
from datetime import datetime
try:
    import docx
except ImportError:
    docx = None

try:
    from src.parsers.base import BaseParser
    from src.core.logger import setup_logger
except ImportError:
    from parsers.base import BaseParser
    from core.logger import setup_logger

logger = setup_logger(__name__)

class DocxParser(BaseParser):
    def parse(self, raw_data: bytes, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Parse Word Document (.docx) bytes."""
        source_name = metadata.get("source", "Unknown Word File")
        timestamp = datetime.now().isoformat()
        try:
            file_like = io.BytesIO(raw_data)
            doc = docx.Document(file_like)
            
            content_parts = []
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
                    
            # Extract tables text
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        content_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(content_parts)
            
            return [{
                "text": full_text,
                "metadata": {
                    "source": source_name,
                    "timestamp": timestamp,
                    "page_number": 1
                }
            }]
        except Exception as e:
            logger.error(f"Error parsing Word Document {source_name}: {e}")
            return []
